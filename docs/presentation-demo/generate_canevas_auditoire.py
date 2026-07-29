# -*- coding: utf-8 -*-
"""Génère le canevas imprimable pour l'auditoire de la démo SIGNELEC / Naissance."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parent / (
    "Canevas_auditoire_demo_signature_electronique_Naissance.docx"
)

GREEN = RGBColor(0x00, 0x6B, 0x31)
DARK = RGBColor(0x1C, 0x2B, 0x2B)
GRAY = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    set_run_font(r, 12, True, GREEN)
    return h


def add_bullet(doc, text):
    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.space_after = Pt(2)
    r = bp.add_run(text)
    set_run_font(r, 10)
    return bp


def style_cell_runs(cell, size=9, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, size, bold, color)


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # En-tête
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("SIFEC"), 22, True, GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(
        p.add_run("Canevas d'orientation — Démonstration"),
        14,
        True,
        DARK,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(
        p.add_run(
            "Intégration de la signature électronique SIGNELEC (GUOT)\n"
            "Module Naissance"
        ),
        11,
        False,
        GRAY,
    )

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(8)
    set_run_font(
        info.add_run(
            "Document remis à l'auditoire  |  À suivre pendant la démo  |  "
            "Durée estimée : 15–20 min"
        ),
        9,
        False,
        GRAY,
    )

    # 1. Objectif
    add_heading_custom(doc, "1. Objectif de la démonstration")
    add_bullet(
        doc,
        "Montrer le fonctionnement de la signature électronique dans SIFEC "
        "après intégration de l'API SIGNELEC.",
    )
    add_bullet(
        doc,
        "Illustrer le parcours métier sur le module Naissance "
        "(déclaration puis acte).",
    )
    add_bullet(
        doc,
        "Mettre en évidence le remplacement de la validation OTP "
        "par le certificat .p12.",
    )

    # 2. Avant / Après
    add_heading_custom(doc, "2. Avant / Après")
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("AVANT", "APRÈS"),
        (
            "Validation par code OTP (SMS)",
            "Signature cryptographique avec certificat .p12",
        ),
        (
            "Preuve limitée à un code temporaire",
            "Empreinte du document + vérification GUOT + cachet institutionnel",
        ),
    ]
    for i, (a, b) in enumerate(rows_data):
        row = table.rows[i]
        row.cells[0].text = a
        row.cells[1].text = b
        for cell in row.cells:
            style_cell_runs(cell, 9, bold=(i == 0), color=GREEN if i == 0 else DARK)

    # 3. Déroulé
    add_heading_custom(doc, "3. Déroulé des étapes (à suivre)")
    steps = [
        (
            "Étape 0 — Introduction",
            "Présentation du contexte SIGNELEC / GUOT et du module Naissance.",
            "~1 min",
        ),
        (
            "Étape 1 — Contexte métier",
            "Ouverture d'une déclaration de naissance déjà arrivée au centre "
            "d'état civil (CEC), en attente de confirmation.",
            "~1 min",
        ),
        (
            "Étape 2 — Signature de la déclaration",
            "L'officier signe électroniquement (fichier .p12 + code PIN).\n"
            "Action écran : « Signer et confirmer ».\n"
            "Effet : dossier confirmé après vérification GUOT.",
            "~3 min",
        ),
        (
            "Étape 3 — Génération de l'acte",
            "Génération de l'acte de naissance.\n"
            "À ce stade, l'acte est créé mais encore en attente de signature "
            "de l'officier.",
            "~2 min",
        ),
        (
            "Étape 4 — Signature de l'acte",
            "Signature électronique de l'acte (même mécanisme .p12).\n"
            "Action écran : « Signer électroniquement ».\n"
            "Effet : acte scellé ; attribution des éléments de preuve.",
            "~3 min",
        ),
        (
            "Étape 5 — Preuve visuelle",
            "Affichage / aperçu du document signé : mention de signature "
            "électronique, date, rôle du signataire.",
            "~2 min",
        ),
        (
            "Étape 6 — Synthèse & questions",
            "Rappel des points clés et échange avec l'auditoire.",
            "~3–5 min",
        ),
    ]

    t2 = doc.add_table(rows=1 + len(steps), cols=3)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t2.rows[0].cells
    hdr[0].text = "Étape"
    hdr[1].text = "Ce qui se passe"
    hdr[2].text = "Durée"
    for cell in hdr:
        style_cell_runs(cell, 9, True, GREEN)

    for i, (title, detail, duree) in enumerate(steps, start=1):
        row = t2.rows[i].cells
        row[0].text = title
        row[1].text = detail
        row[2].text = duree
        style_cell_runs(row[0], 9, True, DARK)
        style_cell_runs(row[1], 9, False, DARK)
        style_cell_runs(row[2], 9, False, DARK)

    # Largeurs approximatives
    for row in t2.rows:
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(1.8)

    # 4. Points à retenir
    add_heading_custom(doc, "4. Points à retenir")
    add_bullet(
        doc,
        "Seuls les postes enrôlés (ex. Officier d'état civil) portent le "
        "certificat — pas les agents de saisie.",
    )
    add_bullet(
        doc,
        "La clé privée reste sur le poste : on signe l'empreinte du document "
        "en local, puis SIGNELEC vérifie.",
    )
    add_bullet(
        doc,
        "Le même modèle s'applique aux autres modules (Mariage, Décès, "
        "Registres, Demande de documents).",
    )
    add_bullet(
        doc,
        "La validation OTP n'est plus utilisée pour ces signatures.",
    )

    # 5. Lexique
    add_heading_custom(doc, "5. Petit lexique")
    lexique = [
        (
            ".p12",
            "Fichier certificat personnel de l'agent signataire "
            "(protégé par un code PIN).",
        ),
        (
            "SIGNELEC / GUOT",
            "Plateforme / API de signature électronique intégrée dans SIFEC.",
        ),
        ("CEC", "Centre d'état civil."),
        (
            "Empreinte (hash)",
            "Identifiant cryptographique du document signé — "
            "toute modification le change.",
        ),
    ]
    for term, defn in lexique:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(1)
        r = bp.add_run(f"{term} — ")
        set_run_font(r, 9, True, GREEN)
        r2 = bp.add_run(defn)
        set_run_font(r2, 9, False, DARK)

    # Pied
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    set_run_font(
        p.add_run(
            "SIFEC  ·  Intégration SIGNELEC  ·  "
            "Support auditoire — Démonstration module Naissance"
        ),
        8,
        False,
        GRAY,
    )

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    main()
