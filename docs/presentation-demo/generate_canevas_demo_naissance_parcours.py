# -*- coding: utf-8 -*-
"""Génère le canevas Word de démonstration SIFEC — parcours naissance (sans signature électronique)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Canevas_demo_parcours_naissance_FS_CEC_Portail.docx"

GREEN = RGBColor(0x00, 0x6B, 0x31)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x4B, 0x55, 0x63)


def set_run_font(run, size=11, bold=False, color=DARK, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_heading_custom(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=GREEN)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True, color=GREEN)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, bullet=False):
    style = "List Bullet" if bullet else None
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + " ")
    set_run_font(r1, size=11, bold=True, color=GREEN)
    r2 = p.add_run(value)
    set_run_font(r2, size=11, color=DARK)
    p.paragraph_format.space_after = Pt(3)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=DARK, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        shade_cell(cell, "006B31")
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            set_cell_text(cell, val, bold=False, size=10)
            if r_idx % 2 == 1:
                shade_cell(cell, "F3F4F6")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_step(doc, num, titre, acteur, duree, objectif, actions, resultat, astuce=None):
    add_subheading(doc, f"Étape {num} — {titre}")
    add_label_value(doc, "Acteur :", acteur)
    add_label_value(doc, "Durée indicative :", duree)
    add_label_value(doc, "Objectif :", objectif)
    p = doc.add_paragraph()
    r = p.add_run("Actions à réaliser")
    set_run_font(r, size=11, bold=True, color=DARK)
    for a in actions:
        add_body(doc, a, bullet=True)
    add_label_value(doc, "Résultat attendu :", resultat)
    if astuce:
        add_label_value(doc, "Astuce démo :", astuce)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # En-tête
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SIFEC")
    set_run_font(r, size=28, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Canevas d’orientation — Démonstration")
    set_run_font(r, size=14, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Parcours complet d’une naissance\n"
        "Formation sanitaire → Centre d’état civil → Portail citoyen → Copie d’acte"
    )
    set_run_font(r, size=12, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("République du Congo  ·  Module Naissance  ·  Démo métier")
    set_run_font(r, size=10, color=MUTED)

    # 1. Objectif
    add_heading_custom(doc, "1. Objectif de la démonstration")
    add_body(
        doc,
        "Montrer, de bout en bout, le circuit métier SIFEC pour une naissance "
        "déclarée en formation sanitaire jusqu’à la délivrance d’une copie d’acte "
        "demandée en ligne.",
    )
    add_body(doc, "Mettre en évidence la continuité numérique entre institutions (FS ↔ CEC).", bullet=True)
    add_body(doc, "Illustrer le rôle de chaque acteur (agent FS, agent/chef CEC, officier, citoyen).", bullet=True)
    add_body(doc, "Démontrer le service au citoyen via le portail (demande de copie + traitement au CEC).", bullet=True)

    # 2. Scénario
    add_heading_custom(doc, "2. Scénario de la démo")
    add_body(
        doc,
        "Un enfant naît à l’Hôpital de Makelekele. La formation sanitaire établit et "
        "transmet le certificat de naissance au Centre d’état civil (Mairie d’arrondissement 1 — Makelekele). "
        "Le CEC valide le dossier, produit la déclaration puis l’acte. L’officier valide/signe l’acte. "
        "Le déclarant retire l’original. Plus tard, une copie est demandée via le portail citoyen, "
        "traitée puis signée au CEC.",
    )

    # 3. Acteurs / comptes
    add_heading_custom(doc, "3. Acteurs et comptes de démonstration")
    add_table(
        doc,
        ["Étape(s)", "Rôle", "Compte (e-mail)", "Institution"],
        [
            ["1", "Agent formation sanitaire", "agentfs@sifec.cg", "Hôpital Makelekele (FS)"],
            ["2–3, 5, 7", "Agent / Chef de service état civil", "agentcec@sifec.cg\ncscec@sifec.cg", "Mairie Arr. 1 Makelekele (CEC)"],
            ["4, 8", "Officier d’état civil", "officiercec@sifec.cg", "Mairie Arr. 1 Makelekele (CEC)"],
            ["6", "Citoyen (déclarant / demandeur)", "Compte portail / usager démo", "Portail citoyen SIFEC"],
        ],
        col_widths=[2.2, 4.2, 4.5, 5.0],
    )
    add_body(
        doc,
        "Mot de passe des comptes seeders démo : 123456 (environnement de démonstration uniquement).",
    )
    add_body(
        doc,
        "Préparer à l’avance : une naissance fictive cohérente (noms, dates, pièces d’identité du déclarant, "
        "lien de parenté), et un acte déjà signé si vous voulez accélérer les étapes 6 à 8.",
    )

    # 4. Séquencement
    add_heading_custom(doc, "4. Séquencement (vue d’ensemble)")
    add_table(
        doc,
        ["N°", "Étape", "Lieu / canal", "Livrable"],
        [
            ["1", "Envoi du certificat de naissance au CEC", "Formation sanitaire", "Certificat transmis au CEC"],
            ["2", "Validation et génération de la déclaration", "CEC", "Déclaration de naissance"],
            ["3", "Génération de l’acte à partir de la déclaration", "CEC", "Acte de naissance généré"],
            ["4", "Signature / validation de l’acte par l’officier", "CEC", "Acte signé, prêt au retrait"],
            ["5", "Enregistrement du retrait par le déclarant", "CEC (guichet)", "Retrait tracé dans SIFEC"],
            ["6", "Demande de copie via le portail", "Portail citoyen", "Demande enregistrée / payée"],
            ["7", "Traitement de la demande par le CEC", "CEC", "Copie préparée"],
            ["8", "Signature de la copie par l’officier", "CEC", "Copie signée délivrable"],
        ],
        col_widths=[1.2, 6.5, 3.8, 4.5],
    )

    # 5. Déroulé
    add_heading_custom(doc, "5. Déroulé des étapes (à suivre)")

    add_step(
        doc,
        1,
        "Envoyer un certificat de naissance au CEC (Formation sanitaire)",
        "Agent formation sanitaire (agentfs@sifec.cg)",
        "3–5 min",
        "Créer le certificat médical de naissance et le transmettre au centre d’état civil partenaire.",
        [
            "Se connecter avec le compte formation sanitaire.",
            "Ouvrir le menu Formation sanitaire / Naissance → Nouveau certificat de naissance.",
            "Saisir l’enfant (noms, sexe, date/heure et lieu de naissance), la mère, le père le cas échéant, et le déclarant.",
            "Joindre les pièces utiles (pièce d’identité, éventuelle attestation) selon le scénario préparé.",
            "Choisir le CEC destinataire (Mairie Arr. 1 Makelekele) puis enregistrer / envoyer le certificat.",
            "Pointer à l’auditoire le numéro / code du certificat généré (à réutiliser aux étapes suivantes).",
        ],
        "Le certificat apparaît côté CEC comme dossier à traiter (en attente de validation / génération de déclaration).",
        "Annoncer clairement le fil rouge : « La FS ne produit pas l’acte — elle ouvre le dossier administratif. »",
    )

    add_step(
        doc,
        2,
        "Valider et générer la déclaration de naissance (CEC)",
        "Agent mairie ou Chef de service état civil (agentcec@ / cscec@sifec.cg)",
        "4–6 min",
        "Contrôler le certificat reçu, compléter si besoin, puis générer la déclaration de naissance.",
        [
            "Se déconnecter / reconnecter avec un compte CEC.",
            "Ouvrir le menu Centre d’état civil / Naissance → Certificats reçus (ou déclarations à générer).",
            "Retrouver le certificat de l’étape 1 (par code, nom de l’enfant ou date).",
            "Vérifier la cohérence des données (identité, parenté, date, institution d’origine).",
            "Valider le dossier puis lancer la génération de la déclaration de naissance.",
            "Montrer l’écran de la déclaration (numéro, parties, statut).",
        ],
        "Une déclaration de naissance est créée et liée au certificat source ; le dossier peut passer à la génération d’acte.",
        "Insister sur la traçabilité : chaque document conserve le lien vers le document précédent.",
    )

    add_step(
        doc,
        3,
        "Générer un acte à partir de la déclaration (CEC)",
        "Agent / Chef de service état civil",
        "2–4 min",
        "Produire l’acte de naissance à partir de la déclaration validée (affectation registre / numéro).",
        [
            "Depuis la déclaration, choisir l’action Générer l’acte de naissance.",
            "Contrôler le registre / la série et le numéro proposés (selon paramétrage du CEC).",
            "Confirmer la génération et afficher l’aperçu de l’acte (maquette).",
            "Indiquer que l’acte est désormais en attente de validation par l’officier d’état civil.",
        ],
        "L’acte existe dans SIFEC avec un numéro ; statut typique : produit / en attente d’approbation de l’officier.",
        "Si un registre de démo est déjà ouvert, la génération est immédiate — vérifier avant la séance.",
    )

    add_step(
        doc,
        4,
        "Signer / valider l’acte par l’officier d’état civil",
        "Officier d’état civil (officiercec@sifec.cg)",
        "2–3 min",
        "Donner force exécutoire à l’acte par la validation de l’officier.",
        [
            "Se connecter avec le compte officier.",
            "Ouvrir la file des actes de naissance à signer / valider.",
            "Ouvrir l’acte généré à l’étape 3, relire les mentions essentielles à voix haute (enfant, parents, n° d’acte).",
            "Lancer la signature / validation de l’acte selon le parcours métier en place sur l’environnement de démo.",
            "Afficher le statut final (acte signé / disponible pour retrait).",
        ],
        "L’acte est signé (ou validé) ; il peut être retiré au guichet par le déclarant.",
        "Sans entrer dans le détail technique : présenter simplement « l’officier authentifie l’acte ». "
        "Cette démo ne couvre pas l’intégration SIGNELEC.",
    )

    add_step(
        doc,
        5,
        "Enregistrement du retrait de l’acte par le déclarant",
        "Agent CEC au guichet + déclarant (présentiel simulé)",
        "2–3 min",
        "Tracer la remise de l’acte original au déclarant.",
        [
            "Menu Retrait / Consultation des actes (ou action Retrait sur l’acte de naissance).",
            "Rechercher l’acte (n° d’acte, nom de l’enfant, code déclaration).",
            "Vérifier l’identité du déclarant (pièce présentée — scénario oral ou pièce déposée en amont).",
            "Enregistrer le retrait (date, agent, éventuellement observations).",
            "Confirmer à l’auditoire que le système conserve la preuve de délivrance.",
        ],
        "Le retrait est historisé ; l’acte original est considéré comme remis.",
        "Bien distinguer original (retrait guichet) et copie (demande ultérieure via portail).",
    )

    add_step(
        doc,
        6,
        "Demande de la copie de l’acte via le portail citoyen",
        "Citoyen / usager (portail)",
        "3–5 min",
        "Montrer le service numérique : le citoyen demande une copie sans retourner remplir un dossier papier.",
        [
            "Ouvrir le portail citoyen SIFEC (session séparée / autre navigateur).",
            "Accéder à Demande de document → Copie d’acte de naissance.",
            "Identifier l’acte (références connues : n° d’acte, noms, date de naissance, CEC de conservation).",
            "Renseigner les coordonnées du demandeur et le motif si demandé.",
            "Valider la demande ; le cas échéant, simuler le paiement associé à la copie.",
            "Noter le code de demande (DD_…) pour le traitement au CEC.",
        ],
        "Une demande de copie est créée et visible côté CEC pour traitement.",
        "Prévoir un usager / une demande déjà partiellement renseignée pour gagner du temps devant l’auditoire.",
    )

    add_step(
        doc,
        7,
        "Traitement de la demande par le CEC",
        "Agent / Chef de service état civil",
        "3–4 min",
        "Instruire la demande portail : contrôle, génération de la copie, préparation à la signature.",
        [
            "Revenir sur le compte CEC.",
            "Ouvrir le module des demandes de documents / copies à traiter.",
            "Retrouver la demande de l’étape 6 (code DD_…).",
            "Vérifier le paiement / le statut, contrôler que l’acte source est bien celui attendu.",
            "Générer / préparer la copie d’acte de naissance (aperçu maquette).",
            "Passer la demande au statut prêt pour signature de l’officier (selon workflow local).",
        ],
        "La copie est produite et liée à la demande ; elle attend la validation de l’officier.",
        "Souligner le gain : le CEC traite une file numérique, pas un dossier papier reconstitutif.",
    )

    add_step(
        doc,
        8,
        "Signer la copie de l’acte de naissance (officier CEC)",
        "Officier d’état civil (officiercec@sifec.cg)",
        "2–3 min",
        "Authentifier la copie pour la rendre délivrable au demandeur.",
        [
            "Se connecter en officier.",
            "Ouvrir les copies / demandes en attente de signature.",
            "Contrôler rapidement les mentions de la copie (identité, n° d’acte, CEC).",
            "Signer / valider la copie.",
            "Montrer le document final (aperçu / impression) et clôturer : la copie peut être retirée ou mise à disposition selon le canal choisi.",
        ],
        "Copie d’acte signée et associée à la demande portail — parcours bout-en-bout terminé.",
        "Conclure en rappelant la chaîne : Certificat FS → Déclaration → Acte → Retrait → Demande portail → Copie signée.",
    )

    # 6. Checklist
    add_heading_custom(doc, "6. Checklist avant la séance")
    add_body(doc, "Environnement SIFEC accessible (back-office + portail) et comptes démo testés.", bullet=True)
    add_body(doc, "Registre de naissance ouvert / disponible pour la génération d’acte au CEC de démo.", bullet=True)
    add_body(doc, "Jeu de données fictives prêt (noms congolais cohérents, dates, pièces).", bullet=True)
    add_body(doc, "Deux navigateurs ou profils (FS/CEC vs portail) pour éviter les mélanges de session.", bullet=True)
    add_body(doc, "Plan B : un acte déjà signé et retiré pour enchaîner directement les étapes 6–8 si le temps presse.", bullet=True)
    add_body(doc, "Chronométrage global conseillé : 25–35 minutes selon le niveau de détail oral.", bullet=True)

    # 7. Messages clés
    add_heading_custom(doc, "7. Messages clés pour l’auditoire")
    add_body(doc, "SIFEC digitalise la chaîne d’état civil, pas seulement un formulaire isolé.", bullet=True)
    add_body(doc, "La formation sanitaire alimente le CEC ; l’officier reste le garant de l’acte.", bullet=True)
    add_body(doc, "Le retrait guichet et la copie portail sont deux services distincts, tous deux tracés.", bullet=True)
    add_body(doc, "Le citoyen gagne du temps : demande de copie en ligne, traitement structuré au CEC.", bullet=True)

    # 8. Lexique
    add_heading_custom(doc, "8. Petit lexique")
    add_label_value(doc, "FS —", "Formation sanitaire (hôpital, centre de santé) qui établit le certificat de naissance.")
    add_label_value(doc, "CEC —", "Centre d’état civil (souvent rattaché à une mairie d’arrondissement).")
    add_label_value(doc, "Certificat de naissance —", "Document source établi à la FS ; ouvre le dossier au CEC.")
    add_label_value(doc, "Déclaration de naissance —", "Formalisation administrative au CEC à partir du certificat.")
    add_label_value(doc, "Acte de naissance —", "Document d’état civil officiel, numéroté, validé par l’officier.")
    add_label_value(doc, "Retrait —", "Remise de l’original au déclarant, enregistrée dans SIFEC.")
    add_label_value(doc, "Copie —", "Reproduction officielle d’un acte déjà établi, souvent demandée via le portail.")
    add_label_value(doc, "Portail citoyen —", "Espace usager pour formuler des demandes de documents (ex. copie).")

    # Pied
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Document d’animation — Démonstration SIFEC · Parcours naissance (métier)")
    set_run_font(r, size=9, color=MUTED)

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
