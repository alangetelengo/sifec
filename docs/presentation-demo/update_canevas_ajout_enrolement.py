# -*- coding: utf-8 -*-
"""Ajoute l'étape d'enrôlement / génération certificat au canevas auditoire."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt

PATH = Path(__file__).resolve().parent / (
    "Canevas_auditoire_demo_signature_electronique_Naissance.docx"
)

NEW_STEPS = [
    (
        "Étape 0 — Introduction",
        "Présentation du contexte SIGNELEC / GUOT et du module Naissance.",
    ),
    (
        "Étape 1 — Génération du certificat électronique",
        "Enrôlement d'un signataire éligible (ex. Officier d'état civil ou chef de service).\n"
        "Action écran : Administration SIGNELEC → Signataires & enrôlements.\n"
        "Effet : génération du certificat personnel .p12 (protégé par PIN), "
        "prérequis à toute signature.",
    ),
    (
        "Étape 2 — Contexte métier",
        "Ouverture d'une déclaration de naissance déjà arrivée au centre "
        "d'état civil (CEC), en attente de confirmation.",
    ),
    (
        "Étape 3 — Signature du certificat de naissance",
        "Le chef de service gestion des malades signe électroniquement "
        "(fichier .p12 + code PIN).\n"
        "Action écran : « Signer et confirmer ».\n"
        "Effet : dossier confirmé après vérification GUOT.",
    ),
    (
        "Étape 4 — Signature de la déclaration",
        "L'officier signe électroniquement (fichier .p12 + code PIN).\n"
        "Action écran : « Signer et confirmer ».\n"
        "Effet : dossier confirmé après vérification GUOT.",
    ),
    (
        "Étape 5 — Génération de l'acte",
        "Génération de l'acte de naissance.\n"
        "À ce stade, l'acte est créé mais encore en attente de signature "
        "de l'officier.",
    ),
    (
        "Étape 6 — Signature de l'acte",
        "Signature électronique de l'acte (même mécanisme .p12).\n"
        "Action écran : « Signer électroniquement ».\n"
        "Effet : acte scellé ; attribution des éléments de preuve.",
    ),
    (
        "Étape 7 — Preuve visuelle",
        "Affichage / aperçu du document signé : mention de signature "
        "électronique, date, rôle du signataire.",
    ),
    (
        "Étape 8 — Synthèse & questions",
        "Rappel des points clés et échange avec l'auditoire.",
    ),
]


def set_cell_text(cell, text, bold=False):
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text
        return

    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)

    p0 = cell.paragraphs[0]
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)

    lines = text.split("\n")
    first = True
    for line in lines:
        if first:
            run = p0.add_run(line)
            first = False
        else:
            run = p0.add_run("\n" + line)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.bold = bold


def main():
    doc = Document(PATH)
    table = doc.tables[1]

    needed = len(NEW_STEPS)
    while len(table.rows) - 1 < needed:
        last_tr = table.rows[-1]._tr
        new_tr = deepcopy(last_tr)
        last_tr.addnext(new_tr)

    for i, (title, detail) in enumerate(NEW_STEPS):
        row = table.rows[i + 1]
        set_cell_text(row.cells[0], title, bold=True)
        set_cell_text(row.cells[1], detail, bold=False)

    doc.save(PATH)
    print(f"OK: {PATH}")
    for i, row in enumerate(doc.tables[1].rows):
        print(f"  {i}: {row.cells[0].text}")


if __name__ == "__main__":
    main()
