from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

# Ensure docs directory exists
os.makedirs('docs', exist_ok=True)

pdf_path = 'docs/GUOT_trust_api_resume.pdf'
docx_path = 'docs/GUOT_trust_api_plan.docx'

pdf_text = [
    ('Objectif', 'Ce guide explique comment intégrer la signature électronique GUOT dans une application métier comme SIFEC, avec un modèle à 3 couches : Layer 1 pour le payload métier, Layer 2 pour le PDF, et Layer 3 pour le cachet institutionnel.'),
    ('Architecture', 'Pattern à 3 couches :\n- Layer 1 : signature du payload métier (/sign)\n- Layer 2 : signature documentaire du PDF (/sign-document)\n- Layer 3 : cachet institutionnel (/seal-document)\nImportant : L2 et L3 se font ensemble sur le même hash PDF, dans la même transaction.'),
    ('Prérequis', 'Variables d’environnement : PKI_TRUST_API_URL, PKI_API_KEY, PKI_CLIENT_CERT, PKI_CLIENT_KEY, PKI_CA_CERT. L’API key est obtenue via la console Signum après validation du tenant GUOT. Pas de GET /status avant chaque signature ; l’appel de signature est l’autorité.'),
    ('Enrôlement PKI', 'Signataires humains : POST /v1/signers, stocker actor_id stable. Institution : institution_id fourni par GUOT/Signum. Ne pas créer le tenant en production. Statut d’acteur : contrôler ponctuellement, pas systématiquement avant chaque signature.'),
    ('Modifications BD', 'Ajouter des colonnes PKI dans toutes les tables de documents officiels, notamment proof_id, payload_hash, actor_id, actor_nom, certificate_ref, signed_at, rfc3161_l1_serial, pdf_content_hash, doc_sig_id, doc_sig_signed_at, rfc3161_l2_serial, doc_seal_id, doc_seal_sealed_at, rfc3161_l3_serial, pdf_path, print_count, first_delivered_at. Ajouter des index pour proof_id, actor_id, doc_sig_id.'),
    ('Layer 1', 'Construire un payload JSON canonique trié alphabétiquement, incluant actor_id, cert_fingerprint, document_id, document_type, transaction_id et les données métier essentielles. Appeler POST /sign avec transaction_id, actor_id, payload, purpose. Stocker proof_id, payload_hash, signed_at, rfc3161_l1_serial.'),
    ('Layer 2', 'Générer le PDF sans cartouche. Calculer le hash SHA-256 du binaire PDF. Appeler POST /sign-document avec document_hash, actor_id, document_type, purpose. Stocker pdf_content_hash, doc_sig_id, doc_sig_signed_at, rfc3161_l2_serial.'),
    ('Layer 3', 'Appeler POST /seal-document avec institution_id, document_hash, purpose, et optionnellement doc_sig_id. Stocker doc_seal_id, doc_seal_sealed_at, rfc3161_l3_serial. Générer le PDF final avec cartouche contenant proof_id, doc_sig_id, doc_seal_id, TSA serials, actor_nom, institution, date et lien de vérification.'),
    ('MinIO', 'Migrer le stockage PDF du disque local vers MinIO S3 compatible. Configurer un disque Laravel minio et installer league/flysystem-aws-s3-v3. Migrer les fichiers existants via une commande Artisan puis basculer FILESYSTEM_DISK=minio.'),
    ('Vérification', 'Vérification publique via portail Signum avec proof_id. Vérification interne via /verify et /verify-document. Vérification hors ligne possible avec OpenSSL et recalcul du hash du payload et du PDF.'),
    ('Cycle de vie des certificats', 'Certificats GUOT valables 1 an. Renouveler à J-30. Statut suspended temporaire, révocation définitive nécessite un nouveau signer. Révocation urgente via POST /v1/signers/{actor_id}/revoke. Suspension tenant renvoie 403 sur les opérations de signature.'),
    ('Checklist production', 'Tenant validé, clé API live configurée, MinIO accessible, mTLS installé si nécessaire, acteurs enrôlés, colonnes PKI ajoutées, QR codes de vérification dans les PDF, tests L1/L2/L3 réussis, téléchargement PDF depuis MinIO, vérification publique et hors ligne validée.'),
]

c = canvas.Canvas(pdf_path, pagesize=letter)
width, height = letter
margin = 72
y = height - margin
c.setFont('Helvetica-Bold', 18)
c.drawString(margin, y, 'Résumé du guide GUOT / trust-api')
y -= 32
c.setFont('Helvetica', 11)
for title, paragraph in pdf_text:
    if y < margin + 120:
        c.showPage()
        y = height - margin
        c.setFont('Helvetica', 11)
    c.setFont('Helvetica-Bold', 13)
    c.drawString(margin, y, title)
    y -= 18
    c.setFont('Helvetica', 10)
    for line in paragraph.replace('\\n', '\n').split('\n'):
        words = line.split(' ')
        line_buf = ''
        for word in words:
            test_line = line_buf + word + ' '
            if c.stringWidth(test_line, 'Helvetica', 10) > width - 2 * margin:
                c.drawString(margin, y, line_buf)
                y -= 14
                line_buf = ''
            line_buf += word + ' '
        if line_buf:
            c.drawString(margin, y, line_buf)
            y -= 14
    y -= 12
c.save()

# Word document
plan = Document()
plan.add_heading('Plan de travail GUOT / SIFEC - 10 jours', level=1)
plan.add_paragraph('Matrice de tâches et responsabilités pour l’intégration GUOT de la signature électronique dans SIFEC.')

rows = [
    ('Jour', 'Tâches principales', 'Responsables'),
    ('J1', 'Analyser l’architecture existante des workflows de signature et recenser les tables de documents concernés.', 'PM, Dev Backend, DBA'),
    ('J2', 'Concevoir et implémenter l’enrôlement PKI des signataires ; stocker actor_id et données de certificat.', 'Dev Backend, PM'),
    ('J3', 'Créer migrations PKI pour les tables de documents et ajouter les index recommandés.', 'DBA, Dev Backend'),
    ('J4', 'Implémenter le client trust-api Laravel et la gestion des erreurs HTTP/API.', 'Dev Backend, DevOps'),
    ('J5', 'Développer la signature Layer 1 (/sign) avec payload canonique et stockage des preuves.', 'Dev Backend, QA'),
    ('J6', 'Développer la signature Layer 2 (/sign-document) et le stockage des métadonnées PDF.', 'Dev Backend, QA'),
    ('J7', 'Développer le cachet Layer 3 (/seal-document) et la génération du PDF final avec cartouche.', 'Dev Backend, QA'),
    ('J8', 'Déployer MinIO, configurer Laravel pour S3, et créer la commande de migration des fichiers existants.', 'DevOps, Dev Backend'),
    ('J9', 'Tester la vérification en ligne et interne, vérifier les PDFs depuis MinIO, valider les liens de vérification.', 'QA, Dev Backend'),
    ('J10', 'Exécuter la checklist de production, préparer la documentation et valider les scénarios de mise en service.', 'PM, QA, DevOps'),
]

table = plan.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Jour'
hdr_cells[1].text = 'Tâches principales'
hdr_cells[2].text = 'Responsables'
for row in rows[1:]:
    cells = table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = row

plan.add_page_break()
plan.add_heading('Détail des responsabilités', level=2)
plan.add_paragraph('PM : pilote le projet, valide les flux, coordonne avec GUOT/Signum et les équipes internes.')
plan.add_paragraph('Dev Backend : implémente les appels trust-api, la logique de signature, les migrations et la génération PDF.')
plan.add_paragraph('DBA : valide les migrations, les index, et assure la cohérence des tables de documents.')
plan.add_paragraph('DevOps : déploie MinIO, configure le stockage S3, installe les dépendances et prépare les environnements.')
plan.add_paragraph('QA : teste les workflows L1/L2/L3, la migration MinIO, la vérification publique et hors ligne.')
plan.save(docx_path)

print('Created', pdf_path)
print('Created', docx_path)
